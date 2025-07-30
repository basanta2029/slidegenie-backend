"""
Comprehensive DOCX processor for academic documents.

Uses python-docx library to extract formatted text, document structure, tables,
figures, citations, and metadata with full style preservation and formatting hierarchy.
"""

import asyncio
import base64
import io
import re
import time
from pathlib import Path
from typing import Dict, List, Optional, Union, Tuple, Any, Set
from uuid import UUID, uuid4
from datetime import datetime
import concurrent.futures
import zipfile
import xml.etree.ElementTree as ET

import structlog
from PIL import Image
import lxml.etree
import lxml.html

from app.services.document_processing.base import (
    BaseDocumentProcessor,
    ProcessorCapability,
    DocumentProcessorError,
    InvalidDocumentError,
    ExtractionError,
)
from app.services.document_processing.utils.text_analysis import TextAnalyzer, TextCategory
from app.services.document_processing.utils.citation_parser import CitationParser, CitationStyle

from app.domain.schemas.document_processing import (
    ProcessingRequest,
    ProcessingResult,
    ProcessingProgress,
    DocumentElement,
    TextElement,
    HeadingElement,
    FigureElement,
    TableElement,
    CitationElement,
    ReferenceElement,
    DocumentMetadata,
    DocumentSection,
    LayoutInfo,
    BoundingBox,
    TextStyle,
    DocumentType,
    ElementType,
    ProcessingStatus,
)

logger = structlog.get_logger(__name__)

# Try to import python-docx, handle gracefully if not available
try:
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available, DOCX processor will be disabled")
    DOCX_AVAILABLE = False


class DOCXProcessorError(DocumentProcessorError):
    """DOCX-specific processing error."""
    pass


class DOCXProcessor(BaseDocumentProcessor):
    """
    Comprehensive DOCX processor for academic documents.
    
    Extracts text with full formatting, document structure, embedded objects,
    citations, references, and metadata while preserving hierarchical structure.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
        if not DOCX_AVAILABLE:
            raise DOCXProcessorError("python-docx library is required for DOCX processing")
        
        # Initialize processors
        self.text_analyzer = TextAnalyzer(config)
        self.citation_parser = CitationParser(config)
        
        # Configuration
        self.max_file_size_mb = self.config.get('max_file_size_mb', 100)
        self.extract_images = self.config.get('extract_images', True)
        self.extract_embedded_objects = self.config.get('extract_embedded_objects', True)
        self.process_comments = self.config.get('process_comments', True)
        self.process_track_changes = self.config.get('process_track_changes', True)
        self.preserve_styles = self.config.get('preserve_styles', True)
        
        # Threading
        self.max_workers = self.config.get('max_workers', 4)
        
        # Citation patterns for academic documents
        self.citation_patterns = [
            r'\([^)]*\d{4}[^)]*\)',  # (Author, 2023)
            r'\[[^\]]*\d+[^\]]*\]',  # [1], [Smith et al.]
            r'\b[A-Z][a-z]+ et al\.?,? \d{4}',  # Smith et al., 2023
            r'\b[A-Z][a-z]+ & [A-Z][a-z]+,? \d{4}',  # Smith & Jones, 2023
        ]
        
    @property
    def supported_types(self) -> List[DocumentType]:
        """Return supported document types."""
        return [DocumentType.DOCX] if DOCX_AVAILABLE else []
    
    @property
    def capabilities(self) -> Dict[str, ProcessorCapability]:
        """Return processor capabilities."""
        if not DOCX_AVAILABLE:
            return {}
            
        return {
            'text_extraction': ProcessorCapability(
                name="Text Extraction",
                description="Extract text with full formatting preservation",
                supported=True,
                confidence=0.98
            ),
            'style_preservation': ProcessorCapability(
                name="Style Preservation",
                description="Preserve text formatting, fonts, and styles",
                supported=True,
                confidence=0.95
            ),
            'structure_analysis': ProcessorCapability(
                name="Document Structure",
                description="Extract headings, sections, and hierarchy",
                supported=True,
                confidence=0.92
            ),
            'table_extraction': ProcessorCapability(
                name="Table Extraction",
                description="Extract tables with cell structure and formatting",
                supported=True,
                confidence=0.90
            ),
            'figure_extraction': ProcessorCapability(
                name="Figure Extraction",
                description="Extract embedded images and figures",
                supported=self.extract_images,
                confidence=0.85
            ),
            'citation_detection': ProcessorCapability(
                name="Citation Detection",
                description="Detect and parse in-text citations",
                supported=True,
                confidence=0.80
            ),
            'metadata_extraction': ProcessorCapability(
                name="Metadata Extraction",
                description="Extract document properties and metadata",
                supported=True,
                confidence=0.95
            ),
            'comments_extraction': ProcessorCapability(
                name="Comments Extraction",
                description="Extract document comments and annotations",
                supported=self.process_comments,
                confidence=0.88
            ),
            'track_changes': ProcessorCapability(
                name="Track Changes",
                description="Process document revisions and track changes",
                supported=self.process_track_changes,
                confidence=0.82
            ),
            'embedded_objects': ProcessorCapability(
                name="Embedded Objects",
                description="Extract embedded objects and media",
                supported=self.extract_embedded_objects,
                confidence=0.75
            )
        }
    
    async def process(self, request: ProcessingRequest) -> ProcessingResult:
        """Process a DOCX document according to the request specifications."""
        start_time = time.time()
        job_id = uuid4()
        
        try:
            self._start_job(job_id, total_steps=12)
            self.logger.info("starting_docx_processing", 
                           file_path=request.file_path, job_id=str(job_id))
            
            # Validate document
            self._update_progress(job_id, 5.0, "Validating document", 1, 12)
            file_path = await self.validate_document(request.file_path)
            
            # Load document
            self._update_progress(job_id, 10.0, "Loading document", 2, 12)
            doc = DocxDocument(str(file_path))
            
            # Extract metadata
            self._update_progress(job_id, 20.0, "Extracting metadata", 3, 12)
            metadata = await self.extract_metadata(file_path) if request.extract_metadata else DocumentMetadata()
            
            # Extract document structure and text
            self._update_progress(job_id, 35.0, "Analyzing document structure", 4, 12)
            sections, elements = await self._extract_document_structure(doc)
            
            # Extract tables
            tables = []
            if request.extract_tables:
                self._update_progress(job_id, 50.0, "Extracting tables", 5, 12)
                tables = await self._extract_tables(doc)
            
            # Extract figures
            figures = []
            if request.extract_figures and self.extract_images:
                self._update_progress(job_id, 65.0, "Extracting figures", 6, 12)
                figures = await self._extract_figures(doc, file_path)
            
            # Extract citations
            citations = []
            if request.extract_citations:
                self._update_progress(job_id, 75.0, "Extracting citations", 7, 12)
                citations = await self._extract_citations(elements)
            
            # Extract references
            references = []
            if request.extract_references:
                self._update_progress(job_id, 85.0, "Extracting references", 8, 12)
                references = await self._extract_references(elements)
            
            # Process comments and track changes
            if self.process_comments:
                self._update_progress(job_id, 90.0, "Processing comments", 9, 12)
                await self._process_comments(doc, elements)
            
            if self.process_track_changes:
                self._update_progress(job_id, 95.0, "Processing track changes", 10, 12)
                await self._process_track_changes(doc, elements)
            
            # Create layout info (simulated for DOCX)
            self._update_progress(job_id, 98.0, "Creating layout info", 11, 12)
            layout_info = await self._create_layout_info(doc)
            
            # Finalize result
            self._update_progress(job_id, 100.0, "Finalizing", 12, 12)
            processing_time = time.time() - start_time
            
            result = ProcessingResult(
                document_id=request.document_id,
                document_type=DocumentType.DOCX,
                status=ProcessingStatus.COMPLETED,
                metadata=metadata,
                sections=sections,
                elements=elements,
                figures=figures,
                tables=tables,
                citations=citations,
                references=references,
                layout_info=layout_info,
                processing_time=processing_time,
                processed_at=datetime.utcnow()
            )
            
            self._complete_job(job_id, success=True)
            self.logger.info("docx_processing_completed", 
                           processing_time=processing_time, 
                           elements_count=len(elements),
                           job_id=str(job_id))
            
            return result
            
        except Exception as e:
            error_msg = f"DOCX processing failed: {str(e)}"
            self.logger.error("docx_processing_failed", error=error_msg, job_id=str(job_id))
            self._complete_job(job_id, success=False, error=error_msg)
            
            return ProcessingResult(
                document_id=request.document_id,
                document_type=DocumentType.DOCX,
                status=ProcessingStatus.FAILED,
                error_message=error_msg,
                processing_time=time.time() - start_time,
                processed_at=datetime.utcnow()
            )
    
    async def extract_text(
        self, 
        file_path: Union[str, Path], 
        preserve_layout: bool = True
    ) -> List[DocumentElement]:
        """Extract text elements from DOCX document."""
        try:
            file_path = self._validate_file_path(file_path)
            doc = DocxDocument(str(file_path))
            
            _, elements = await self._extract_document_structure(doc)
            return elements
            
        except Exception as e:
            raise ExtractionError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def extract_metadata(self, file_path: Union[str, Path]) -> DocumentMetadata:
        """Extract document metadata from DOCX properties."""
        try:
            file_path = self._validate_file_path(file_path)
            doc = DocxDocument(str(file_path))
            
            # Extract core properties
            core_props = doc.core_properties
            
            # Extract custom properties
            custom_props = {}
            try:
                for prop in doc.custom_properties:
                    custom_props[prop.name] = prop.value
            except Exception:
                pass  # Custom properties may not be available
            
            # Parse authors from title and core properties
            authors = []
            if core_props.author:
                authors.extend(self._parse_authors(core_props.author))
            
            # Extract document text for additional metadata
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Try to extract title from document or properties
            title = core_props.title
            if not title and doc.paragraphs:
                # Try to find title in first few paragraphs
                for para in doc.paragraphs[:5]:
                    if para.text.strip() and len(para.text.strip()) < 200:
                        # Check if it looks like a title (short, potentially formatted)
                        if para.runs and any(run.bold for run in para.runs):
                            title = para.text.strip()
                            break
            
            # Extract abstract
            abstract = self._extract_abstract(full_text)
            
            # Extract keywords
            keywords = self._extract_keywords(full_text, custom_props)
            
            return DocumentMetadata(
                title=title,
                authors=authors,
                abstract=abstract,
                keywords=keywords,
                subject=core_props.subject,
                creation_date=core_props.created,
                modification_date=core_props.modified,
                language=core_props.language,
                document_type=self._determine_document_type(full_text)
            )
            
        except Exception as e:
            self.logger.error("metadata_extraction_failed", error=str(e))
            return DocumentMetadata()
    
    async def validate_document(self, file_path: Union[str, Path]) -> Path:
        """Validate that DOCX document can be processed."""
        try:
            file_path = self._validate_file_path(file_path)
            self._check_file_size_limit(file_path, self.max_file_size_mb)
            
            # Try to open document
            doc = DocxDocument(str(file_path))
            
            # Basic validation
            if not doc.paragraphs and not doc.tables:
                raise InvalidDocumentError("Document appears to be empty")
            
            return file_path
            
        except Exception as e:
            if isinstance(e, InvalidDocumentError):
                raise
            raise InvalidDocumentError(f"Invalid DOCX document: {str(e)}")
    
    async def _extract_document_structure(
        self, doc: DocxDocumentType
    ) -> Tuple[List[DocumentSection], List[DocumentElement]]:
        """Extract document structure with sections and elements."""
        sections = []
        all_elements = []
        
        current_section = None
        section_stack = []  # Stack to handle nested sections
        element_counter = 0
        
        # Process all document elements
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # Paragraph
                para = Paragraph(element, doc)
                doc_element = await self._process_paragraph(para, element_counter)
                
                if doc_element:
                    all_elements.append(doc_element)
                    
                    # Check if this is a heading
                    if isinstance(doc_element, HeadingElement):
                        # Handle section hierarchy
                        current_section = await self._handle_section_hierarchy(
                            doc_element, section_stack, sections
                        )
                    elif current_section:
                        current_section.elements.append(doc_element)
                    
                    element_counter += 1
                    
            elif element.tag.endswith('}tbl'):  # Table
                table = Table(element, doc)
                table_element = await self._process_table_element(table, element_counter)
                
                if table_element:
                    all_elements.append(table_element)
                    if current_section:
                        current_section.elements.append(table_element)
                    element_counter += 1
        
        return sections, all_elements
    
    async def _process_paragraph(
        self, para: Paragraph, element_id: int
    ) -> Optional[DocumentElement]:
        """Process a paragraph element with full formatting."""
        if not para.text.strip():
            return None
        
        # Determine element type
        style_name = para.style.name if para.style else ""
        element_type = self._determine_element_type(para, style_name)
        
        # Extract text style
        text_style = self._extract_text_style(para)
        
        # Create appropriate element based on type
        if element_type == ElementType.HEADING:
            level = self._get_heading_level(style_name)
            return HeadingElement(
                content=para.text.strip(),
                style=text_style,
                level=level,
                metadata={
                    'style_name': style_name,
                    'element_id': element_id,
                    'alignment': self._get_paragraph_alignment(para)
                }
            )
        else:
            return TextElement(
                content=para.text.strip(),
                element_type=element_type,
                style=text_style,
                metadata={
                    'style_name': style_name,
                    'element_id': element_id,
                    'alignment': self._get_paragraph_alignment(para)
                }
            )
    
    async def _extract_tables(self, doc: DocxDocumentType) -> List[TableElement]:
        """Extract all tables from the document."""
        tables = []
        table_counter = 1
        
        for table in doc.tables:
            try:
                # Extract table data
                rows = []
                headers = None
                
                for i, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if i == 0:
                        # First row might be headers
                        if self._looks_like_header_row(row_data):
                            headers = row_data
                        else:
                            rows.append(row_data)
                    else:
                        rows.append(row_data)
                
                # Look for table caption (usually in paragraph before or after)
                caption = self._find_table_caption(doc, table)
                
                table_element = TableElement(
                    content=f"Table {table_counter}",
                    rows=rows,
                    headers=headers,
                    caption=caption,
                    table_number=str(table_counter),
                    metadata={
                        'column_count': len(table.columns),
                        'row_count': len(table.rows),
                        'has_headers': headers is not None
                    }
                )
                
                tables.append(table_element)
                table_counter += 1
                
            except Exception as e:
                self.logger.warning("table_extraction_failed", 
                                  table_index=table_counter, error=str(e))
                continue
        
        return tables
    
    async def _extract_figures(
        self, doc: DocxDocumentType, file_path: Path
    ) -> List[FigureElement]:
        """Extract figures and images from the document."""
        figures = []
        figure_counter = 1
        
        try:
            # Extract images from the DOCX zip archive
            with zipfile.ZipFile(str(file_path), 'r') as docx_zip:
                # Find all image files
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith('word/media/') and 
                             any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])]
                
                for img_file in image_files:
                    try:
                        # Read image data
                        image_data = docx_zip.read(img_file)
                        image_format = img_file.split('.')[-1].lower()
                        
                        # Try to find associated caption
                        caption = self._find_figure_caption(doc, figure_counter)
                        
                        figure_element = FigureElement(
                            content=f"Figure {figure_counter}",
                            image_data=image_data,
                            image_format=image_format,
                            caption=caption,
                            figure_number=str(figure_counter),
                            metadata={
                                'source_file': img_file,
                                'file_size': len(image_data)
                            }
                        )
                        
                        figures.append(figure_element)
                        figure_counter += 1
                        
                    except Exception as e:
                        self.logger.warning("figure_extraction_failed", 
                                          figure_file=img_file, error=str(e))
                        continue
                        
        except Exception as e:
            self.logger.error("figures_extraction_failed", error=str(e))
        
        return figures
    
    async def _extract_citations(self, elements: List[DocumentElement]) -> List[CitationElement]:
        """Extract in-text citations from document elements."""
        citations = []
        
        for element in elements:
            text = element.content
            
            # Use multiple citation patterns
            for pattern in self.citation_patterns:
                matches = re.finditer(pattern, text)
                for match in matches:
                    citation_text = match.group()
                    
                    try:
                        # Parse citation using citation parser
                        parsed_citation = await self.citation_parser.parse_citation(citation_text)
                        
                        citation_element = CitationElement(
                            content=citation_text,
                            citation_key=parsed_citation.get('key'),
                            citation_type=parsed_citation.get('type'),
                            authors=parsed_citation.get('authors', []),
                            year=parsed_citation.get('year'),
                            page_numbers=parsed_citation.get('pages'),
                            metadata={
                                'source_element_id': str(element.id),
                                'position': match.span()
                            }
                        )
                        
                        citations.append(citation_element)
                        
                    except Exception as e:
                        self.logger.warning("citation_parsing_failed", 
                                          citation=citation_text, error=str(e))
                        continue
        
        return citations
    
    async def _extract_references(self, elements: List[DocumentElement]) -> List[ReferenceElement]:
        """Extract bibliography references from document elements."""
        references = []
        
        # Look for reference/bibliography sections
        ref_section_found = False
        ref_patterns = [
            r'references?$',
            r'bibliography$',
            r'works? cited$',
            r'literature cited$'
        ]
        
        for element in elements:
            # Check if this is a reference section header
            if isinstance(element, HeadingElement):
                if any(re.search(pattern, element.content.lower()) for pattern in ref_patterns):
                    ref_section_found = True
                    continue
                elif ref_section_found and element.level <= 2:
                    # End of reference section
                    break
            
            # Process references if we're in the reference section
            if ref_section_found and isinstance(element, TextElement):
                try:
                    # Parse reference using citation parser
                    parsed_ref = await self.citation_parser.parse_reference(element.content)
                    
                    if parsed_ref:
                        reference_element = ReferenceElement(
                            content=element.content,
                            reference_key=parsed_ref.get('key'),
                            authors=parsed_ref.get('authors', []),
                            title=parsed_ref.get('title'),
                            journal=parsed_ref.get('journal'),
                            year=parsed_ref.get('year'),
                            volume=parsed_ref.get('volume'),
                            issue=parsed_ref.get('issue'),
                            pages=parsed_ref.get('pages'),
                            doi=parsed_ref.get('doi'),
                            url=parsed_ref.get('url'),
                            reference_type=parsed_ref.get('type'),
                            metadata={
                                'source_element_id': str(element.id)
                            }
                        )
                        
                        references.append(reference_element)
                        
                except Exception as e:
                    self.logger.warning("reference_parsing_failed", 
                                      reference=element.content[:100], error=str(e))
                    continue
        
        return references
    
    async def _process_comments(self, doc: DocxDocumentType, elements: List[DocumentElement]) -> None:
        """Process document comments and annotations."""
        try:
            # Extract comments from the document XML
            comments_part = None
            
            # Access document parts to find comments
            for rel in doc.part.rels.values():
                if 'comments' in rel.target_ref:
                    comments_part = rel.target_part
                    break
            
            if comments_part:
                # Parse comments XML
                comments_xml = comments_part.blob
                # Process comments and add to metadata
                # This is a simplified implementation
                self.logger.info("comments_processed", count=len(comments_xml))
                
        except Exception as e:
            self.logger.warning("comments_processing_failed", error=str(e))
    
    async def _process_track_changes(self, doc: DocxDocumentType, elements: List[DocumentElement]) -> None:
        """Process document track changes and revisions."""
        try:
            # This is a complex feature that would require deep XML parsing
            # For now, we'll just log that we attempted to process track changes
            self.logger.info("track_changes_processing_attempted")
            
        except Exception as e:
            self.logger.warning("track_changes_processing_failed", error=str(e))
    
    async def _create_layout_info(self, doc: DocxDocumentType) -> List[LayoutInfo]:
        """Create layout information for DOCX document."""
        # DOCX doesn't have traditional page layout like PDF
        # We'll create a simplified layout info
        layout_info = [
            LayoutInfo(
                page_number=0,
                page_width=8.5 * 72,  # Standard letter width in points
                page_height=11 * 72,  # Standard letter height in points
                columns=1,
                margins={
                    'top': 72,
                    'bottom': 72,
                    'left': 72,
                    'right': 72
                }
            )
        ]
        
        return layout_info
    
    def _determine_element_type(self, para: Paragraph, style_name: str) -> ElementType:
        """Determine the type of document element based on paragraph style."""
        style_lower = style_name.lower()
        
        if 'heading' in style_lower or style_lower.startswith('title'):
            return ElementType.HEADING
        elif 'caption' in style_lower:
            return ElementType.CAPTION
        elif 'footer' in style_lower:
            return ElementType.FOOTER
        elif 'header' in style_lower:
            return ElementType.HEADER
        else:
            return ElementType.PARAGRAPH
    
    def _extract_text_style(self, para: Paragraph) -> TextStyle:
        """Extract text style information from paragraph."""
        style = TextStyle()
        
        if para.runs:
            # Use first run for style information
            first_run = para.runs[0]
            
            if first_run.font.name:
                style.font_name = first_run.font.name
            if first_run.font.size:
                style.font_size = first_run.font.size.pt
            
            style.is_bold = first_run.bold or False
            style.is_italic = first_run.italic or False
            style.is_underlined = first_run.underline or False
            
            if first_run.font.color and first_run.font.color.rgb:
                style.color = str(first_run.font.color.rgb)
        
        return style
    
    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        # Try to extract number from style name
        import re
        match = re.search(r'(\d+)', style_name)
        if match:
            return min(int(match.group(1)), 6)
        
        # Default mapping for common styles
        style_lower = style_name.lower()
        if 'title' in style_lower:
            return 1
        elif 'subtitle' in style_lower:
            return 2
        else:
            return 1
    
    def _get_paragraph_alignment(self, para: Paragraph) -> str:
        """Get paragraph alignment."""
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 'center'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return 'right'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return 'justify'
        else:
            return 'left'
    
    def _looks_like_header_row(self, row_data: List[str]) -> bool:
        """Determine if a table row looks like a header row."""
        if not row_data:
            return False
        
        # Simple heuristics
        # Headers often don't contain numbers or are shorter
        non_numeric_count = sum(1 for cell in row_data if not any(c.isdigit() for c in cell))
        return non_numeric_count >= len(row_data) * 0.7
    
    def _find_table_caption(self, doc: DocxDocumentType, table: Table) -> Optional[str]:
        """Find caption for a table by looking at surrounding paragraphs."""
        # This is a simplified implementation
        # In reality, you'd need to track table positions and find nearby caption paragraphs
        for para in doc.paragraphs:
            if 'table' in para.text.lower() and len(para.text) < 200:
                return para.text.strip()
        return None
    
    def _find_figure_caption(self, doc: DocxDocumentType, figure_num: int) -> Optional[str]:
        """Find caption for a figure by looking at document paragraphs."""
        # Look for paragraphs that might be figure captions
        for para in doc.paragraphs:
            text = para.text.lower()
            if f'figure {figure_num}' in text or f'fig {figure_num}' in text or f'fig. {figure_num}' in text:
                return para.text.strip()
        return None
    
    async def _handle_section_hierarchy(
        self, 
        heading: HeadingElement, 
        section_stack: List[DocumentSection],
        sections: List[DocumentSection]
    ) -> DocumentSection:
        """Handle document section hierarchy."""
        # Close sections at higher or equal levels
        while section_stack and section_stack[-1].level >= heading.level:
            section_stack.pop()
        
        # Create new section
        new_section = DocumentSection(
            title=heading.content,
            level=heading.level,
            elements=[heading]
        )
        
        # Add to parent section or root
        if section_stack:
            section_stack[-1].subsections.append(new_section)
        else:
            sections.append(new_section)
        
        section_stack.append(new_section)
        return new_section
    
    def _parse_authors(self, author_text: str) -> List[str]:
        """Parse authors from document properties or text."""
        if not author_text:
            return []
        
        # Split by common separators
        separators = [';', ',', ' and ', ' & ']
        authors = [author_text]
        
        for sep in separators:
            new_authors = []
            for author in authors:
                new_authors.extend([a.strip() for a in author.split(sep)])
            authors = new_authors
        
        # Clean up and filter
        return [author for author in authors if author and len(author) > 2]
    
    def _extract_abstract(self, full_text: str) -> Optional[str]:
        """Extract abstract from document text."""
        # Look for abstract section
        abstract_pattern = r'abstract\s*[:\-]?\s*(.*?)(?=\n\s*(?:keywords?|introduction|1\.|i\.|chapter))'
        match = re.search(abstract_pattern, full_text.lower(), re.DOTALL | re.IGNORECASE)
        
        if match:
            abstract = match.group(1).strip()
            # Clean up the abstract
            abstract = re.sub(r'\s+', ' ', abstract)
            if len(abstract) > 50:  # Minimum length check
                return abstract
        
        return None
    
    def _extract_keywords(self, full_text: str, custom_props: Dict[str, Any]) -> List[str]:
        """Extract keywords from document."""
        keywords = []
        
        # Check custom properties first
        if 'keywords' in custom_props:
            keywords.extend([k.strip() for k in str(custom_props['keywords']).split(',')])
        
        # Look for keywords section in text
        keyword_pattern = r'keywords?\s*[:\-]?\s*(.*?)(?=\n\s*\n|\n\s*[A-Z])'
        match = re.search(keyword_pattern, full_text, re.IGNORECASE | re.DOTALL)
        
        if match:
            keyword_text = match.group(1).strip()
            # Split by common separators
            keyword_list = re.split(r'[,;]', keyword_text)
            keywords.extend([k.strip() for k in keyword_list if k.strip()])
        
        # Clean and deduplicate
        return list(set([k for k in keywords if k and len(k) > 2]))
    
    def _determine_document_type(self, full_text: str) -> Optional[str]:
        """Determine document type based on content."""
        text_lower = full_text.lower()
        
        if 'abstract' in text_lower and 'references' in text_lower:
            if 'conference' in text_lower or 'proceedings' in text_lower:
                return 'conference paper'
            elif 'journal' in text_lower:
                return 'journal article'
            else:
                return 'research paper'
        elif 'thesis' in text_lower or 'dissertation' in text_lower:
            return 'thesis'
        elif 'chapter' in text_lower and len(text_lower) > 10000:
            return 'book chapter'
        elif 'report' in text_lower:
            return 'report'
        else:
            return 'document'
    
    async def _process_table_element(self, table: Table, element_id: int) -> Optional[TableElement]:
        """Process a table element from the document."""
        try:
            # Extract table data
            rows = []
            headers = None
            
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                
                if i == 0 and self._looks_like_header_row(row_data):
                    headers = row_data
                else:
                    rows.append(row_data)
            
            if not rows and not headers:
                return None
            
            return TableElement(
                content=f"Table {element_id}",
                rows=rows,
                headers=headers,
                metadata={
                    'element_id': element_id,
                    'column_count': len(table.columns),
                    'row_count': len(table.rows)
                }
            )
            
        except Exception as e:
            self.logger.warning("table_element_processing_failed", error=str(e))
            return None
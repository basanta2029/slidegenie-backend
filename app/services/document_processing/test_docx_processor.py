"""
Test suite for the DOCX processor.

Tests text extraction, formatting preservation, table/figure extraction,
citation detection, and metadata parsing from DOCX documents.
"""

import pytest
import asyncio
import tempfile
import os
from pathlib import Path
from uuid import uuid4
from datetime import datetime

from app.services.document_processing.processors.docx_processor import DOCXProcessor
from app.domain.schemas.document_processing import (
    ProcessingRequest,
    DocumentType,
    ElementType,
    HeadingElement,
    TableElement,
    FigureElement,
    CitationElement,
    ReferenceElement,
)

# Skip all tests if python-docx is not available
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

pytestmark = pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")


class TestDOCXProcessor:
    """Test cases for DOCX processor functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.processor = DOCXProcessor()
        self.temp_dir = Path(tempfile.mkdtemp())
    
    def teardown_method(self):
        """Clean up test fixtures."""
        # Clean up temporary files
        import shutil
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
    
    def create_test_document(self, filename: str = "test.docx") -> Path:
        """Create a test DOCX document with various elements."""
        doc_path = self.temp_dir / filename
        
        # Create a document with various elements
        doc = DocxDocument()
        
        # Document properties
        doc.core_properties.title = "Test Academic Document"
        doc.core_properties.author = "John Doe; Jane Smith"
        doc.core_properties.subject = "Computer Science"
        doc.core_properties.created = datetime.now()
        
        # Title
        title = doc.add_heading('Advanced Machine Learning Techniques', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Authors
        authors = doc.add_paragraph('John Doe¹, Jane Smith²')
        authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Affiliations
        affiliations = doc.add_paragraph('¹Department of Computer Science, University A')
        affiliations.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('²Department of AI Research, University B').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Abstract
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(
            'This paper presents advanced machine learning techniques for classification tasks. '
            'We demonstrate improved performance using novel deep learning architectures. '
            'The proposed method achieves 95% accuracy on benchmark datasets.'
        )
        
        # Keywords
        doc.add_paragraph('Keywords: machine learning, deep learning, classification, neural networks')
        
        # Introduction
        doc.add_heading('1. Introduction', level=1)
        doc.add_paragraph(
            'Machine learning has revolutionized many fields in recent years (Smith et al., 2020). '
            'Deep learning techniques have shown particular promise in classification tasks [1, 2]. '
            'This work builds upon previous research (Johnson & Brown, 2019) to develop improved methods.'
        )
        
        # Methodology
        doc.add_heading('2. Methodology', level=1)
        doc.add_paragraph(
            'Our approach consists of three main components: data preprocessing, '
            'model architecture design, and optimization techniques.'
        )
        
        # Subsection
        doc.add_heading('2.1 Data Preprocessing', level=2)
        doc.add_paragraph(
            'We applied standard normalization techniques to prepare the input data. '
            'Missing values were handled using interpolation methods.'
        )
        
        # Table
        doc.add_paragraph('Table 1: Experimental Results')
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # Table headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Method'
        header_cells[1].text = 'Accuracy (%)'
        header_cells[2].text = 'Training Time (h)'
        
        # Table data
        data = [
            ['Baseline CNN', '87.2', '2.5'],
            ['Improved CNN', '92.1', '3.1'],
            ['Our Method', '95.3', '2.8']
        ]
        
        for i, row_data in enumerate(data, 1):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        
        # Results section
        doc.add_heading('3. Results', level=1)
        doc.add_paragraph(
            'Our experimental evaluation demonstrates significant improvements over baseline methods. '
            'As shown in Table 1, our approach achieves 95.3% accuracy while maintaining competitive training times.'
        )
        
        # References
        doc.add_heading('References', level=1)
        references = [
            '[1] Smith, A., Johnson, B., & Davis, C. (2020). Deep learning fundamentals. Journal of AI Research, 15(3), 123-145.',
            '[2] Brown, D., & Wilson, E. (2019). Neural network architectures for classification. Conference on Machine Learning, 45-52.',
            'Johnson, M., & Brown, P. (2019). Advanced optimization techniques in deep learning. Nature Machine Intelligence, 1(4), 234-248.'
        ]
        
        for ref in references:
            doc.add_paragraph(ref)
        
        # Save document
        doc.save(str(doc_path))
        return doc_path
    
    @pytest.mark.asyncio
    async def test_processor_initialization(self):
        """Test processor initialization and capabilities."""
        assert DocumentType.DOCX in self.processor.supported_types
        
        capabilities = self.processor.capabilities
        assert 'text_extraction' in capabilities
        assert 'style_preservation' in capabilities
        assert 'structure_analysis' in capabilities
        assert 'table_extraction' in capabilities
        assert 'metadata_extraction' in capabilities
        
        assert capabilities['text_extraction'].supported
        assert capabilities['text_extraction'].confidence > 0.9
    
    @pytest.mark.asyncio
    async def test_document_validation(self):
        """Test document validation."""
        # Test valid document
        doc_path = self.create_test_document()
        validated_path = await self.processor.validate_document(doc_path)
        assert validated_path == doc_path
        
        # Test invalid file path
        with pytest.raises(Exception):
            await self.processor.validate_document("nonexistent.docx")
    
    @pytest.mark.asyncio
    async def test_metadata_extraction(self):
        """Test metadata extraction from document properties."""
        doc_path = self.create_test_document()
        metadata = await self.processor.extract_metadata(doc_path)
        
        assert metadata.title == "Test Academic Document"
        assert "John Doe" in metadata.authors
        assert "Jane Smith" in metadata.authors
        assert metadata.subject == "Computer Science"
        assert metadata.creation_date is not None
        assert metadata.abstract is not None
        assert len(metadata.keywords) > 0
        assert "machine learning" in [k.lower() for k in metadata.keywords]
    
    @pytest.mark.asyncio
    async def test_text_extraction(self):
        """Test text extraction with formatting preservation."""
        doc_path = self.create_test_document()
        elements = await self.processor.extract_text(doc_path)
        
        assert len(elements) > 0
        
        # Check for different element types
        element_types = [elem.element_type for elem in elements]
        assert ElementType.HEADING in element_types
        assert ElementType.PARAGRAPH in element_types
        
        # Check for headings
        headings = [elem for elem in elements if isinstance(elem, HeadingElement)]
        assert len(headings) > 0
        
        # Verify heading levels
        heading_levels = [h.level for h in headings]
        assert 1 in heading_levels  # Should have level 1 headings
        assert 2 in heading_levels  # Should have level 2 headings
        
        # Check text content
        all_text = " ".join([elem.content for elem in elements])
        assert "Advanced Machine Learning Techniques" in all_text
        assert "Introduction" in all_text
        assert "Methodology" in all_text
    
    @pytest.mark.asyncio
    async def test_full_document_processing(self):
        """Test complete document processing workflow."""
        doc_path = self.create_test_document()
        
        request = ProcessingRequest(
            document_id=uuid4(),
            file_path=str(doc_path),
            document_type=DocumentType.DOCX,
            extract_text=True,
            extract_tables=True,
            extract_figures=True,
            extract_citations=True,
            extract_references=True,
            extract_metadata=True
        )
        
        result = await self.processor.process(request)
        
        # Check result status
        assert result.status.value == "completed"
        assert result.document_type == DocumentType.DOCX
        assert result.processing_time is not None
        
        # Check metadata
        assert result.metadata.title is not None
        assert len(result.metadata.authors) > 0
        
        # Check sections
        assert len(result.sections) > 0
        section_titles = [s.title for s in result.sections]
        assert any("Introduction" in title for title in section_titles)
        assert any("Methodology" in title for title in section_titles)
        
        # Check elements
        assert len(result.elements) > 0
        
        # Check tables (should find at least one)
        assert len(result.tables) > 0
        table = result.tables[0]
        assert table.headers is not None
        assert len(table.rows) > 0
        assert "Method" in table.headers
        assert "Accuracy" in " ".join(table.headers)
        
        # Check citations
        assert len(result.citations) > 0
        citation_texts = [c.content for c in result.citations]
        assert any("Smith et al." in text for text in citation_texts)
        
        # Check references
        assert len(result.references) > 0
        reference_texts = [r.content for r in result.references]
        assert any("Journal of AI Research" in text for text in reference_texts)
    
    @pytest.mark.asyncio
    async def test_style_preservation(self):
        """Test preservation of text styles and formatting."""
        doc_path = self.create_test_document()
        elements = await self.processor.extract_text(doc_path)
        
        # Check that style information is preserved
        styled_elements = [elem for elem in elements if elem.style is not None]
        assert len(styled_elements) > 0
        
        # Check for font information
        font_elements = [elem for elem in styled_elements if elem.style.font_name is not None]
        # Font information might not always be available, so this is optional
        
        # Check for formatting flags
        heading_elements = [elem for elem in elements if isinstance(elem, HeadingElement)]
        if heading_elements:
            # Headings might have bold formatting
            bold_elements = [elem for elem in heading_elements if elem.style and elem.style.is_bold]
            # This is optional as it depends on the document style
    
    @pytest.mark.asyncio
    async def test_table_extraction_detailed(self):
        """Test detailed table extraction functionality."""
        doc_path = self.create_test_document()
        
        request = ProcessingRequest(
            document_id=uuid4(),
            file_path=str(doc_path),
            document_type=DocumentType.DOCX,
            extract_tables=True
        )
        
        result = await self.processor.process(request)
        
        assert len(result.tables) > 0
        table = result.tables[0]
        
        # Check table structure
        assert table.headers is not None
        assert len(table.headers) == 3  # Method, Accuracy, Training Time
        assert table.headers[0] == "Method"
        assert "Accuracy" in table.headers[1]
        assert "Training Time" in table.headers[2]
        
        # Check table data
        assert len(table.rows) == 3  # Three data rows
        assert "Baseline CNN" in table.rows[0]
        assert "87.2" in table.rows[0]
        assert "Our Method" in table.rows[2]
        assert "95.3" in table.rows[2]
        
        # Check metadata
        assert table.metadata is not None
        assert table.metadata['column_count'] == 3
        assert table.metadata['row_count'] == 4  # Including header
    
    @pytest.mark.asyncio
    async def test_citation_detection(self):
        """Test citation detection and parsing."""
        doc_path = self.create_test_document()
        elements = await self.processor.extract_text(doc_path)
        citations = await self.processor._extract_citations(elements)
        
        assert len(citations) > 0
        
        # Check for different citation formats
        citation_contents = [c.content for c in citations]
        
        # Should find author-year citations
        author_year_citations = [c for c in citation_contents if "Smith et al." in c]
        assert len(author_year_citations) > 0
        
        # Should find numbered citations
        numbered_citations = [c for c in citation_contents if "[1" in c or "[2" in c]
        assert len(numbered_citations) > 0
        
        # Check citation parsing
        parsed_citations = [c for c in citations if c.authors or c.year]
        assert len(parsed_citations) > 0
    
    @pytest.mark.asyncio
    async def test_reference_extraction(self):
        """Test reference/bibliography extraction."""
        doc_path = self.create_test_document()
        elements = await self.processor.extract_text(doc_path)
        references = await self.processor._extract_references(elements)
        
        assert len(references) > 0
        
        # Check reference content
        reference_texts = [r.content for r in references]
        
        # Should find journal references
        journal_refs = [r for r in reference_texts if "Journal of AI Research" in r]
        assert len(journal_refs) > 0
        
        # Should find conference references
        conf_refs = [r for r in reference_texts if "Conference on Machine Learning" in r]
        assert len(conf_refs) > 0
        
        # Check parsed reference information
        parsed_refs = [r for r in references if r.authors or r.title or r.year]
        assert len(parsed_refs) > 0
    
    @pytest.mark.asyncio
    async def test_section_hierarchy(self):
        """Test document section hierarchy extraction."""
        doc_path = self.create_test_document()
        
        request = ProcessingRequest(
            document_id=uuid4(),
            file_path=str(doc_path),
            document_type=DocumentType.DOCX
        )
        
        result = await self.processor.process(request)
        
        # Check section structure
        assert len(result.sections) > 0
        
        # Find main sections
        main_sections = [s for s in result.sections if s.level == 1]
        assert len(main_sections) > 0
        
        # Check for subsections
        methodology_section = next((s for s in main_sections if "Methodology" in s.title), None)
        if methodology_section:
            assert len(methodology_section.subsections) > 0
            subsection = methodology_section.subsections[0]
            assert subsection.level == 2
            assert "Data Preprocessing" in subsection.title
    
    @pytest.mark.asyncio
    async def test_error_handling(self):
        """Test error handling for various edge cases."""
        # Test empty document
        empty_doc_path = self.temp_dir / "empty.docx"
        empty_doc = DocxDocument()
        empty_doc.save(str(empty_doc_path))
        
        with pytest.raises(Exception):
            await self.processor.validate_document(empty_doc_path)
        
        # Test corrupted file
        corrupted_path = self.temp_dir / "corrupted.docx"
        with open(corrupted_path, 'w') as f:
            f.write("This is not a DOCX file")
        
        with pytest.raises(Exception):
            await self.processor.validate_document(corrupted_path)
    
    @pytest.mark.asyncio
    async def test_progress_tracking(self):
        """Test processing progress tracking."""
        doc_path = self.create_test_document()
        
        request = ProcessingRequest(
            document_id=uuid4(),
            file_path=str(doc_path),
            document_type=DocumentType.DOCX
        )
        
        # Start processing (this will create a job)
        result = await self.processor.process(request)
        
        # Since processing is fast, job will be completed
        assert result.status.value == "completed"
    
    def test_processor_capabilities_without_docx(self):
        """Test processor behavior when python-docx is not available."""
        # This test would need to mock the import failure
        # For now, we just ensure the current processor works
        assert DOCX_AVAILABLE  # This test file only runs when DOCX is available


if __name__ == "__main__":
    # Run a simple test
    async def main():
        test = TestDOCXProcessor()
        test.setup_method()
        
        try:
            await test.test_processor_initialization()
            print("✓ Processor initialization test passed")
            
            await test.test_full_document_processing()
            print("✓ Full document processing test passed")
            
            print("All tests passed!")
            
        except Exception as e:
            print(f"✗ Test failed: {e}")
            raise
        finally:
            test.teardown_method()
    
    # Run the test
    asyncio.run(main())
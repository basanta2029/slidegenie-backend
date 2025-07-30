"""
Example usage of the DOCX processor for academic documents.

Demonstrates various features including text extraction, style preservation,
table/figure extraction, citation detection, and metadata parsing.
"""

import asyncio
import json
from pathlib import Path
from uuid import uuid4
from typing import Dict, Any

from app.services.document_processing.processors.docx_processor import DOCXProcessor
from app.domain.schemas.document_processing import (
    ProcessingRequest, 
    DocumentType, 
    ElementType
)


async def demonstrate_docx_processing():
    """Demonstrate comprehensive DOCX processing capabilities."""
    
    print("=== DOCX Processor Demonstration ===\n")
    
    # Initialize processor
    config = {
        'max_file_size_mb': 50,
        'extract_images': True,
        'extract_embedded_objects': True,
        'process_comments': True,
        'process_track_changes': True,
        'preserve_styles': True
    }
    
    processor = DOCXProcessor(config)
    
    # Display processor capabilities
    print("1. Processor Capabilities:")
    capabilities = processor.capabilities
    for name, capability in capabilities.items():
        status = "✓" if capability.supported else "✗"
        print(f"   {status} {capability.name}: {capability.description}")
        print(f"     Confidence: {capability.confidence:.1%}")
    print()
    
    # Note: This example assumes you have a DOCX file to process
    # For demonstration, we'll create a sample document programmatically
    
    try:
        from docx import Document
        sample_doc_path = await create_sample_document()
        print(f"2. Created sample document: {sample_doc_path}\n")
        
        # Process the document
        await process_document_example(processor, sample_doc_path)
        
    except ImportError:
        print("2. python-docx not available - using mock examples\n")
        await demonstrate_mock_processing(processor)


async def create_sample_document() -> Path:
    """Create a sample academic DOCX document for demonstration."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Machine Learning in Healthcare Applications"
    doc.core_properties.author = "Dr. Sarah Johnson; Prof. Michael Chen"
    doc.core_properties.subject = "Artificial Intelligence"
    doc.core_properties.keywords = "machine learning, healthcare, AI, medical diagnosis"
    
    # Title
    title = doc.add_heading('Machine Learning in Healthcare Applications', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Authors and affiliations
    authors = doc.add_paragraph('Dr. Sarah Johnson¹, Prof. Michael Chen²')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('¹Department of Computer Science, Medical University').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('²AI Research Institute, Tech University').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This study investigates the application of machine learning techniques in healthcare "
        "diagnostics. We analyze the effectiveness of deep learning models for medical image "
        "analysis and patient outcome prediction. Our results demonstrate significant "
        "improvements in diagnostic accuracy, achieving 94.2% precision in disease detection. "
        "The proposed framework integrates multiple ML algorithms to provide comprehensive "
        "healthcare analytics solutions."
    )
    doc.add_paragraph(abstract_text)
    
    # Keywords
    doc.add_paragraph('Keywords: machine learning, healthcare, medical diagnosis, deep learning, AI')
    
    # Introduction
    doc.add_heading('1. Introduction', level=1)
    intro_text = (
        "The integration of artificial intelligence in healthcare has shown remarkable progress "
        "in recent years (Smith & Johnson, 2023). Machine learning algorithms have demonstrated "
        "exceptional capabilities in medical image analysis [1, 2]. Previous studies have shown "
        "that deep learning models can achieve accuracy levels comparable to expert radiologists "
        "(Chen et al., 2022). This research builds upon existing work to develop more robust "
        "and interpretable ML systems for clinical applications."
    )
    doc.add_paragraph(intro_text)
    
    # Literature Review
    doc.add_heading('2. Literature Review', level=1)
    doc.add_paragraph(
        "Recent advances in computer vision have enabled breakthrough applications in medical "
        "imaging. Convolutional Neural Networks (CNNs) have been particularly successful in "
        "radiological image analysis (Brown et al., 2021). Natural Language Processing techniques "
        "have also been applied to electronic health records with promising results [3]."
    )
    
    # Methodology
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        "Our approach consists of three main components: data preprocessing, model development, "
        "and evaluation metrics. We employed a multi-modal approach combining image data with "
        "clinical parameters."
    )
    
    # Data preprocessing subsection
    doc.add_heading('3.1 Data Preprocessing', level=2)
    doc.add_paragraph(
        "Medical images were normalized and augmented using standard techniques. "
        "Patient data was anonymized and preprocessed according to HIPAA guidelines."
    )
    
    # Model architecture subsection
    doc.add_heading('3.2 Model Architecture', level=2)
    doc.add_paragraph(
        "We implemented a hybrid architecture combining ResNet-50 for image feature extraction "
        "with LSTM networks for temporal sequence modeling."
    )
    
    # Results table
    doc.add_heading('4. Results', level=1)
    doc.add_paragraph('Table 1: Performance Comparison of Different ML Models')
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # Headers
    headers = ['Model', 'Accuracy (%)', 'Precision (%)', 'Recall (%)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Data
    data = [
        ['Traditional CNN', '87.3', '85.1', '89.2'],
        ['ResNet-50', '91.7', '90.4', '92.8'],
        ['Our Hybrid Model', '94.2', '93.8', '94.6'],
        ['Ensemble Method', '95.1', '94.7', '95.4']
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    # Discussion
    doc.add_heading('5. Discussion', level=1)
    doc.add_paragraph(
        "The results demonstrate the effectiveness of our hybrid approach. "
        "The integration of multiple data modalities significantly improved diagnostic accuracy. "
        "Statistical analysis using t-tests confirmed the significance of our improvements (p < 0.001)."
    )
    
    # Limitations
    doc.add_heading('5.1 Limitations', level=2)
    doc.add_paragraph(
        "This study has several limitations including dataset size and potential bias in "
        "patient selection. Future work should address these concerns with larger, "
        "more diverse datasets."
    )
    
    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "We presented a novel machine learning framework for healthcare applications "
        "that achieves state-of-the-art performance in diagnostic tasks. The proposed "
        "system shows great promise for clinical deployment and could significantly "
        "improve patient outcomes."
    )
    
    # References
    doc.add_heading('References', level=1)
    references = [
        "[1] Anderson, K., Lee, S., & Wilson, R. (2023). Deep learning for medical image analysis: A comprehensive review. Nature Medicine, 29(4), 123-145. doi:10.1038/s41591-023-02246-3",
        "[2] Zhang, L., Kumar, P., & Thompson, M. (2022). Automated diagnosis using convolutional neural networks. Journal of Medical AI, 15(2), 78-92.",
        "[3] Rodriguez, C., & Patel, N. (2021). NLP applications in electronic health records. Healthcare Informatics Review, 8(1), 34-47.",
        "Brown, A., Davis, J., & Miller, T. (2021). Computer vision in radiology: Current applications and future prospects. Radiology Today, 45(3), 156-168.",
        "Chen, W., Liu, Y., & Garcia, M. (2022). AI-assisted medical diagnosis: Performance evaluation and clinical implementation. Medical AI Quarterly, 12(4), 201-215.",
        "Smith, D., & Johnson, H. (2023). Artificial intelligence in healthcare: Opportunities and challenges. New England Journal of Medicine, 388(15), 1234-1245."
    ]
    
    for ref in references:
        doc.add_paragraph(ref)
    
    # Save document
    doc_path = Path("sample_healthcare_ml_paper.docx")
    doc.save(str(doc_path))
    
    return doc_path


async def process_document_example(processor: DOCXProcessor, doc_path: Path):
    """Demonstrate document processing with detailed output."""
    
    print("3. Document Processing Example:")
    print(f"   Processing: {doc_path}")
    
    # Create processing request
    request = ProcessingRequest(
        document_id=uuid4(),
        file_path=str(doc_path),
        document_type=DocumentType.DOCX,
        extract_text=True,
        extract_tables=True,
        extract_figures=True,
        extract_citations=True,
        extract_references=True,
        extract_metadata=True,
        preserve_layout=True
    )
    
    # Process document
    result = await processor.process(request)
    
    # Display results
    print(f"   Status: {result.status}")
    print(f"   Processing time: {result.processing_time:.2f}s")
    print(f"   Elements extracted: {len(result.elements)}")
    print()
    
    # Show metadata
    print("4. Document Metadata:")
    metadata = result.metadata
    print(f"   Title: {metadata.title}")
    print(f"   Authors: {', '.join(metadata.authors)}")
    print(f"   Subject: {metadata.subject}")
    print(f"   Keywords: {', '.join(metadata.keywords)}")
    print(f"   Document type: {metadata.document_type}")
    if metadata.abstract:
        print(f"   Abstract: {metadata.abstract[:100]}...")
    print()
    
    # Show document structure
    print("5. Document Structure:")
    for i, section in enumerate(result.sections, 1):
        indent = "  " * section.level
        print(f"   {indent}{section.level}. {section.title}")
        print(f"   {indent}   Elements: {len(section.elements)}")
        if section.subsections:
            for subsection in section.subsections:
                sub_indent = "  " * subsection.level
                print(f"   {sub_indent}{subsection.level}. {subsection.title}")
    print()
    
    # Show element types
    print("6. Element Analysis:")
    element_types = {}
    for element in result.elements:
        elem_type = element.element_type
        element_types[elem_type] = element_types.get(elem_type, 0) + 1
    
    for elem_type, count in element_types.items():
        print(f"   {elem_type}: {count}")
    print()
    
    # Show headings
    print("7. Document Headings:")
    headings = [elem for elem in result.elements if elem.element_type == ElementType.HEADING]
    for heading in headings[:10]:  # Show first 10 headings
        level_indent = "  " * getattr(heading, 'level', 1)
        print(f"   {level_indent}Level {getattr(heading, 'level', 1)}: {heading.content}")
    print()
    
    # Show tables
    print("8. Tables Extracted:")
    for i, table in enumerate(result.tables, 1):
        print(f"   Table {i}: {table.table_number or f'#{i}'}")
        if table.caption:
            print(f"   Caption: {table.caption}")
        if table.headers:
            print(f"   Headers: {', '.join(table.headers)}")
        print(f"   Dimensions: {len(table.rows)} rows × {len(table.headers or table.rows[0] if table.rows else [])} columns")
        
        # Show first few rows
        if table.rows:
            print("   Sample data:")
            for row in table.rows[:3]:
                print(f"     {' | '.join(str(cell)[:20] for cell in row)}")
        print()
    
    # Show citations
    print("9. Citations Found:")
    for i, citation in enumerate(result.citations[:10], 1):  # Show first 10
        print(f"   Citation {i}: {citation.content}")
        if citation.authors:
            print(f"     Authors: {', '.join(citation.authors)}")
        if citation.year:
            print(f"     Year: {citation.year}")
        print()
    
    # Show references
    print("10. References Extracted:")
    for i, reference in enumerate(result.references[:5], 1):  # Show first 5
        print(f"   Reference {i}:")
        print(f"     {reference.content[:100]}...")
        if reference.authors:
            print(f"     Authors: {', '.join(reference.authors)}")
        if reference.title:
            print(f"     Title: {reference.title[:50]}...")
        if reference.journal:
            print(f"     Journal: {reference.journal}")
        if reference.year:
            print(f"     Year: {reference.year}")
        print()
    
    # Show style information
    print("11. Style Analysis:")
    styled_elements = [elem for elem in result.elements if elem.style is not None]
    print(f"   Elements with style info: {len(styled_elements)}")
    
    if styled_elements:
        # Analyze fonts
        fonts = set()
        bold_count = italic_count = 0
        
        for elem in styled_elements:
            if elem.style.font_name:
                fonts.add(elem.style.font_name)
            if elem.style.is_bold:
                bold_count += 1
            if elem.style.is_italic:
                italic_count += 1
        
        if fonts:
            print(f"   Fonts used: {', '.join(fonts)}")
        print(f"   Bold elements: {bold_count}")
        print(f"   Italic elements: {italic_count}")
    print()
    
    print("Processing completed successfully! ✓")


async def demonstrate_mock_processing(processor: DOCXProcessor):
    """Demonstrate processor capabilities without actual document."""
    
    print("3. Mock Processing Example (python-docx not available):")
    print("   This would demonstrate:")
    print("   - Text extraction with style preservation")
    print("   - Document structure analysis (headings, sections)")
    print("   - Table extraction with formatting")
    print("   - Figure and image extraction")
    print("   - Citation detection and parsing")
    print("   - Reference extraction from bibliography")
    print("   - Metadata extraction from document properties")
    print("   - Comment and track changes processing")
    print()
    
    print("4. Expected Output Structure:")
    mock_structure = {
        "metadata": {
            "title": "Academic Paper Title",
            "authors": ["Author 1", "Author 2"],
            "keywords": ["keyword1", "keyword2"],
            "abstract": "Paper abstract text...",
            "document_type": "research paper"
        },
        "sections": [
            {"title": "Introduction", "level": 1, "elements": 5},
            {"title": "Literature Review", "level": 1, "elements": 8},
            {"title": "Methodology", "level": 1, "elements": 12},
            {"title": "Results", "level": 1, "elements": 6},
            {"title": "Discussion", "level": 1, "elements": 4},
            {"title": "Conclusion", "level": 1, "elements": 2}
        ],
        "statistics": {
            "total_elements": 37,
            "headings": 6,
            "paragraphs": 28,
            "tables": 2,
            "figures": 1,
            "citations": 15,
            "references": 25
        }
    }
    
    print(json.dumps(mock_structure, indent=2))


async def demonstrate_advanced_features():
    """Demonstrate advanced DOCX processing features."""
    
    print("=== Advanced Features ===\n")
    
    print("1. Style Preservation:")
    print("   ✓ Font family and size extraction")
    print("   ✓ Bold, italic, underline formatting")
    print("   ✓ Text color and highlighting")
    print("   ✓ Paragraph alignment and spacing")
    print()
    
    print("2. Document Structure Analysis:")
    print("   ✓ Hierarchical heading detection")
    print("   ✓ Section and subsection organization")
    print("   ✓ Automatic section numbering")
    print("   ✓ Cross-reference resolution")
    print()
    
    print("3. Table Processing:")
    print("   ✓ Cell content extraction")
    print("   ✓ Header row detection")
    print("   ✓ Table formatting preservation")
    print("   ✓ Merged cell handling")
    print("   ✓ Caption and numbering")
    print()
    
    print("4. Figure and Media Extraction:")
    print("   ✓ Embedded image extraction")
    print("   ✓ Image format detection")
    print("   ✓ Caption association")
    print("   ✓ Figure numbering and references")
    print()
    
    print("5. Citation and Reference Processing:")
    print("   ✓ Multiple citation formats (APA, MLA, IEEE)")
    print("   ✓ In-text citation detection")
    print("   ✓ Bibliography parsing")
    print("   ✓ Author and publication info extraction")
    print("   ✓ DOI and URL detection")
    print()
    
    print("6. Advanced Document Features:")
    print("   ✓ Comments and annotations")
    print("   ✓ Track changes processing")  
    print("   ✓ Embedded objects handling")
    print("   ✓ Custom document properties")
    print("   ✓ Field and formula extraction")
    print()
    
    print("7. Error Handling and Robustness:")
    print("   ✓ Corrupted document detection")
    print("   ✓ Missing dependency graceful handling")
    print("   ✓ Large document processing")
    print("   ✓ Memory-efficient streaming")
    print("   ✓ Progress tracking and cancellation")


if __name__ == "__main__":
    print("DOCX Processor - Comprehensive Academic Document Processing\n")
    
    async def main():
        await demonstrate_docx_processing()
        print("\n" + "="*60 + "\n")
        await demonstrate_advanced_features()
    
    # Run the demonstration
    asyncio.run(main())